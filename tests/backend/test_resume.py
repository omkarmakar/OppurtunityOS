"""Resume parsing endpoint tests."""

from __future__ import annotations

import os
import tempfile
import uuid

from docx import Document
from fastapi.testclient import TestClient

SAMPLE_TEX = r"""\documentclass{article}
\usepackage{geometry}
\begin{document}

\section{Skills}
Python, Java, \textbf{Deep Learning}

\section{Experience}
Senior Developer at Acme
2020 -- 2024
\begin{itemize}
    \item Built APIs with FastAPI
\end{itemize}

\section{Education}
MIT
2014 -- 2018
B.S. in Computer Science

\end{document}
"""


class TestResumeParse:
    def test_parse_docx_returns_200(self, client: TestClient) -> None:
        doc = Document()
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("Python, Go")
        path = os.path.join(tempfile.gettempdir(), f"test_{uuid.uuid4().hex}.docx")
        doc.save(path)
        with open(path, "rb") as f:
            resp = client.post("/api/v1/resume/parse", files={"file": ("resume.docx", f, "application/octet-stream")})
        os.unlink(path)
        assert resp.status_code == 200

    def test_parse_returns_all_sections(self, client: TestClient) -> None:
        doc = Document()
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("Python, FastAPI, Docker")
        doc.add_paragraph("EXPERIENCE")
        doc.add_paragraph("Dev at Acme")
        doc.add_paragraph("2020 - 2024")
        doc.add_paragraph("- Built APIs")
        doc.add_paragraph("EDUCATION")
        doc.add_paragraph("MIT")
        doc.add_paragraph("2014 - 2018")
        doc.add_paragraph("B.S. CS")
        doc.add_paragraph("PROJECTS")
        doc.add_paragraph("My App")
        doc.add_paragraph("- cool feature")
        doc.add_paragraph("github.com/user/app")
        path = os.path.join(tempfile.gettempdir(), f"test_{uuid.uuid4().hex}.docx")
        doc.save(path)
        with open(path, "rb") as f:
            resp = client.post("/api/v1/resume/parse", files={"file": ("resume.docx", f, "application/octet-stream")})
        os.unlink(path)
        data = resp.json()
        assert len(data["skills"]) >= 3
        assert len(data["education"]) >= 1
        assert len(data["experience"]) >= 1
        assert len(data["projects"]) >= 1
        assert data["file_name"] == "resume.docx"

    def test_parse_tex_returns_200(self, client: TestClient) -> None:
        path = os.path.join(tempfile.gettempdir(), f"test_{uuid.uuid4().hex}.tex")
        with open(path, "w", encoding="utf-8") as f:
            f.write(SAMPLE_TEX)
        with open(path, "rb") as f:
            resp = client.post("/api/v1/resume/parse", files={"file": ("resume.tex", f, "application/octet-stream")})
        os.unlink(path)
        assert resp.status_code == 200
        data = resp.json()
        assert "Python" in data["skills"]
        assert "Deep Learning" in data["skills"]

    def test_parse_tex_returns_all_sections(self, client: TestClient) -> None:
        path = os.path.join(tempfile.gettempdir(), f"test_{uuid.uuid4().hex}.tex")
        with open(path, "w", encoding="utf-8") as f:
            f.write(SAMPLE_TEX)
        with open(path, "rb") as f:
            resp = client.post("/api/v1/resume/parse", files={"file": ("resume.tex", f, "application/octet-stream")})
        os.unlink(path)
        assert resp.status_code == 200
        data = resp.json()
        assert len(data["skills"]) >= 2
        assert len(data["education"]) >= 1
        assert len(data["experience"]) >= 1
        assert data["file_name"] == "resume.tex"

    def test_parse_invalid_file_returns_400(self, client: TestClient) -> None:
        resp = client.post("/api/v1/resume/parse", files={"file": ("test.txt", b"not a resume", "text/plain")})
        assert resp.status_code == 400

    def test_parse_and_save_updates_profile(self, client: TestClient) -> None:
        uid = str(uuid.uuid4())
        client.post("/api/v1/profiles", json={"user_id": uid, "display_name": "Parser Test"})

        doc = Document()
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("Rust, Go")
        path = os.path.join(tempfile.gettempdir(), f"test_{uuid.uuid4().hex}.docx")
        doc.save(path)
        with open(path, "rb") as f:
            resp = client.post(
                f"/api/v1/resume/parse-and-save/{uid}",
                files={"file": ("resume.docx", f, "application/octet-stream")},
            )
        os.unlink(path)
        assert resp.status_code == 200
        assert "Rust" in resp.json()["skills"]

    def test_parse_and_save_no_profile_returns_404(self, client: TestClient) -> None:
        uid = str(uuid.uuid4())
        doc = Document()
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("X")
        path = os.path.join(tempfile.gettempdir(), f"test_{uuid.uuid4().hex}.docx")
        doc.save(path)
        with open(path, "rb") as f:
            resp = client.post(
                f"/api/v1/resume/parse-and-save/{uid}",
                files={"file": ("resume.docx", f, "application/octet-stream")},
            )
        os.unlink(path)
        assert resp.status_code == 404
